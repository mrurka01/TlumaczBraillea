from datetime import datetime
import PyPDF2
from docx import Document as DocxDocument
from odf import text, teletype
from odf.opendocument import load as odf_wczytaj
from striprtf.striprtf import rtf_to_text
import os
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import pypandoc


class ImportExportPolskiBraille:
    OBSLUGIWANE_IMPORTY = {'.txt', '.pdf', '.docx', '.odt', '.rtf'}
    OBSLUGIWANE_EKSPORTY = {'.brf', '.bse', '.bes', '.txt', '.pdf'}

    def __init__(self, menadzer):
        self.menadzer = menadzer

    def importuj_plik(self, sciezka_pliku: str, kodowanie: str = 'utf-8') -> tuple[bool, str]:
        rozszerzenie = Path(sciezka_pliku).suffix.lower()
        if rozszerzenie not in self.OBSLUGIWANE_IMPORTY:
            return False, f"Format {rozszerzenie} nie jest obsługiwany dla Polski → Braille"
        return self.menadzer.importuj_plik(sciezka_pliku, kodowanie)

    def eksportuj_plik(self, sciezka_docelowa: str) -> tuple[bool, str]:
        rozszerzenie = Path(sciezka_docelowa).suffix.lower()
        if rozszerzenie not in self.OBSLUGIWANE_EKSPORTY:
            return False, f"Format {rozszerzenie} nie jest obsługiwany dla Polski → Braille"

        zawartosc = self.menadzer.pobierz_zawartosc()
        if zawartosc is None:
            return False, "Brak zawartości do eksportu"

        try:
            if rozszerzenie in ['.brf', '.bse', '.bes', '.txt']:
                with open(sciezka_docelowa, 'w', encoding='utf-8') as f:
                    f.write(zawartosc)
                return True, None

            if rozszerzenie == '.pdf':
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                mozliwe_sciezki = [
                    os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'),
                    os.path.join(os.path.dirname(__file__), 'DejaVuSans.ttf'),
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                    'C:\\Windows\\Fonts\\DejaVuSans.ttf'
                ]
                font_path = None
                for p in mozliwe_sciezki:
                    if os.path.exists(p):
                        font_path = p
                        break
                if font_path is None:
                    return False, "Brak czcionki DejaVuSans.ttf. Umieść ją w ./fonts/DejaVuSans.ttf"

                try:
                    pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))
                except Exception:
                    pass

                c = canvas.Canvas(sciezka_docelowa, pagesize=A4)
                szer, wys = A4
                margin_x = 50
                y = wys - 50
                line_height = 18
                font_size = 14

                c.setFont("DejaVuSans", font_size)

                for linia in zawartosc.splitlines():
                    max_chars = int((szer - 2 * margin_x) / (font_size * 0.6))
                    while len(linia) > max_chars:
                        fragment = linia[:max_chars]
                        c.drawString(margin_x, y, fragment)
                        linia = linia[max_chars:]
                        y -= line_height
                        if y < 50:
                            c.showPage()
                            c.setFont("DejaVuSans", font_size)
                            y = wys - 50
                    c.drawString(margin_x, y, linia)
                    y -= line_height
                    if y < 50:
                        c.showPage()
                        c.setFont("DejaVuSans", font_size)
                        y = wys - 50
                c.save()
                return True, None

            return False, f"Nieobsługiwane rozszerzenie: {rozszerzenie}"

        except Exception as e:
            return False, f"Błąd eksportu (Polski→Braille): {e}"


class ImportExportBraillePolski:
    OBSLUGIWANE_IMPORTY = {
        '.txt', '.pdf', '.docx', '.odt', '.rtf',
        '.brf', '.bse', '.bes', '.jpg', '.jpeg', '.png'
    }
    OBSLUGIWANE_EKSPORTY = {'.txt', '.pdf', '.docx', '.odt', '.rtf'}

    def __init__(self, menadzer):
        self.menadzer = menadzer

    def importuj_plik(self, sciezka_pliku: str, kodowanie: str = 'utf-8') -> tuple[bool, str]:
        rozszerzenie = Path(sciezka_pliku).suffix.lower()
        if rozszerzenie not in self.OBSLUGIWANE_IMPORTY:
            return False, f"Format {rozszerzenie} nie jest obsługiwany dla Braille → Polski"

        if rozszerzenie in ['.jpg', '.jpeg', '.png']:
            return False, "Obrazy powinny być obsłużone przez OCR; wywołaj _importuj_obraz_braille() w main"

        if rozszerzenie in ['.brf', '.bse', '.bes']:
            try:
                with open(sciezka_pliku, 'r', encoding=kodowanie, errors='ignore') as f:
                    zawartosc = f.read()
                if not zawartosc.strip():
                    return False, "Plik Braille jest pusty"
                self.menadzer.utworz_nowy_dokument(zawartosc)
                return True, None
            except Exception as e:
                return False, f"Błąd odczytu pliku Braille: {e}"

        return self.menadzer.importuj_plik(sciezka_pliku, kodowanie)

    def eksportuj_plik(self, sciezka_docelowa: str) -> tuple[bool, str]:
        rozszerzenie = Path(sciezka_docelowa).suffix.lower()
        if rozszerzenie not in self.OBSLUGIWANE_EKSPORTY:
            return False, f"Format {rozszerzenie} nie jest obsługiwany dla Braille → Polski"

        zawartosc = self.menadzer.pobierz_zawartosc()
        if zawartosc is None:
            return False, "Brak zawartości do eksportu"

        try:
            if rozszerzenie == '.txt':
                with open(sciezka_docelowa, 'w', encoding='utf-8') as f:
                    f.write(zawartosc)
                return True, None

            if rozszerzenie == '.pdf':
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                mozliwe_sciezki = [
                    os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'),
                    os.path.join(os.path.dirname(__file__), 'DejaVuSans.ttf'),
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                    'C:\\Windows\\Fonts\\DejaVuSans.ttf'
                ]
                font_path = None
                for p in mozliwe_sciezki:
                    if os.path.exists(p):
                        font_path = p
                        break
                if font_path:
                    try:
                        pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))
                        chosen_font = "DejaVuSans"
                    except Exception:
                        chosen_font = "Helvetica"
                else:
                    chosen_font = "Helvetica"

                c = canvas.Canvas(sciezka_docelowa, pagesize=A4)
                szer, wys = A4
                margin_x = 50
                y = wys - 50
                line_height = 16
                font_size = 12

                c.setFont(chosen_font, font_size)

                for linia in zawartosc.splitlines():
                    max_chars = int((szer - 2 * margin_x) / (font_size * 0.6))
                    while len(linia) > max_chars:
                        fragment = linia[:max_chars]
                        c.drawString(margin_x, y, fragment)
                        linia = linia[max_chars:]
                        y -= line_height
                        if y < 50:
                            c.showPage()
                            c.setFont(chosen_font, font_size)
                            y = wys - 50
                    c.drawString(margin_x, y, linia)
                    y -= line_height
                    if y < 50:
                        c.showPage()
                        c.setFont(chosen_font, font_size)
                        y = wys - 50
                c.save()
                return True, None

            if rozszerzenie == '.docx':
                from docx import Document
                doc = Document()
                for linia in zawartosc.splitlines():
                    doc.add_paragraph(linia)
                doc.save(sciezka_docelowa)
                return True, None

            if rozszerzenie == '.odt':
                from odf.opendocument import OpenDocumentText
                from odf.text import P
                doc = OpenDocumentText()
                for linia in zawartosc.splitlines():
                    doc.text.addElement(P(text=linia))
                doc.save(sciezka_docelowa)
                return True, None

            if rozszerzenie == '.rtf':
                try:
                    import pypandoc
                    pypandoc.convert_text(zawartosc, 'rtf', format='md', outputfile=sciezka_docelowa, extra_args=['--standalone'])
                    return True, None
                except Exception:
                    with open(sciezka_docelowa, 'w', encoding='utf-8') as f:
                        f.write(zawartosc)
                    return True, None

            return False, f"Nieobsługiwane rozszerzenie: {rozszerzenie}"

        except Exception as e:
            return False, f"Błąd eksportu (Braille→Polski): {e}"

class Dokument:
    def __init__(self, zawartosc: str = "", metadane: dict = None):
        self.zawartosc = zawartosc
        self.metadane = metadane or {}
        self.metadane.setdefault('utworzono', datetime.now())
        self.metadane.setdefault('zmodyfikowano', datetime.now())
        self.metadane.setdefault('oryginalny_format', None)
        self.metadane.setdefault('sciezka', None)
        self._historia = []
        self._indeks_historii = -1

    def pobierz_zawartosc(self) -> str:
        return self.zawartosc

    def ustaw_zawartosc(self, zawartosc: str, zapisz_historie: bool = True):
        if zapisz_historie and self.zawartosc != zawartosc:
            self._historia = self._historia[:self._indeks_historii + 1]
            self._historia.append(self.zawartosc)
            self._indeks_historii += 1
            if len(self._historia) > 50:
                self._historia.pop(0)
                self._indeks_historii -= 1

        self.zawartosc = zawartosc
        self.metadane['zmodyfikowano'] = datetime.now()

    def cofnij(self) -> bool:
        if self._indeks_historii >= 0:
            self.zawartosc = self._historia[self._indeks_historii]
            self._indeks_historii -= 1
            self.metadane['zmodyfikowano'] = datetime.now()
            return True
        return False

    def ponow(self) -> bool:
        if self._indeks_historii < len(self._historia) - 1:
            self._indeks_historii += 1
            self.zawartosc = self._historia[self._indeks_historii]
            self.metadane['zmodyfikowano'] = datetime.now()
            return True
        return False

    def pobierz_statystyki(self) -> dict:
        return {
            'znaki': len(self.zawartosc),
            'znaki_bez_spacji': len(self.zawartosc.replace(' ', '')),
            'slowa': len(self.zawartosc.split()),
            'linie': self.zawartosc.count('\n') + 1,
            'akapity': len([p for p in self.zawartosc.split('\n\n') if p.strip()])
        }


class MenadzerDokumentow:
    OBSLUGIWANE_FORMATY = {'.txt', '.pdf', '.docx', '.odt', '.rtf', '.brf', '.bse', '.bes'}
    OBSLUGIWANE_EKSPORTY = {'.txt', '.pdf', '.docx', '.odt', '.rtf', '.brf'}

    def __init__(self):
        self.aktualny_dokument: Dokument = None

    def importuj_plik(self, sciezka_pliku: str, kodowanie: str = 'utf-8') -> tuple[bool, str]:
        if not os.path.exists(sciezka_pliku):
            return False, "Plik nie istnieje"

        rozszerzenie = Path(sciezka_pliku).suffix.lower()

        if rozszerzenie not in self.OBSLUGIWANE_FORMATY:
            return False, f"Format {rozszerzenie} nie jest obsługiwany"

        try:
            metody = {
                '.txt': self._importuj_txt,
                '.pdf': self._importuj_pdf,
                '.docx': self._importuj_docx,
                '.odt': self._importuj_odt,
                '.rtf': self._importuj_rtf,
                '.brf': self._importuj_braille,
                '.bse': self._importuj_braille,
                '.bes': self._importuj_braille,
            }

            sukces, zawartosc, blad = metody[rozszerzenie](sciezka_pliku, kodowanie)
            if not sukces:
                return False, blad

            self.aktualny_dokument = Dokument(
                zawartosc=zawartosc,
                metadane={
                    'oryginalny_format': rozszerzenie,
                    'sciezka': sciezka_pliku,
                    'nazwa_pliku': os.path.basename(sciezka_pliku)
                }
            )

            return True, None

        except Exception as e:
            return False, f"Błąd podczas importu: {str(e)}"

    def _importuj_txt(self, sciezka, kodowanie):
        try:
            with open(sciezka, 'r', encoding=kodowanie) as f:
                zawartosc = f.read()
                if not zawartosc:
                    return False, "", "Plik jest pusty"
                return True, zawartosc, None
        except Exception as e:
            return False, "", f"Błąd TXT: {str(e)}"

    def _importuj_pdf(self, sciezka, _=None):
        try:
            zawartosc = []
            with open(sciezka, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                for strona in pdf.pages:
                    tekst = strona.extract_text()
                    if tekst:
                        zawartosc.append(tekst)
            pelna = '\n\n'.join(zawartosc)
            if not pelna.strip():
                return False, "", "PDF nie zawiera tekstu"
            return True, pelna, None
        except Exception as e:
            return False, "", f"Błąd PDF: {str(e)}"

    def _importuj_docx(self, sciezka, _=None):
        try:
            dokument = DocxDocument(sciezka)
            teksty = [p.text for p in dokument.paragraphs if p.text.strip()]
            return True, '\n'.join(teksty), None
        except Exception as e:
            return False, "", f"Błąd DOCX: {str(e)}"

    def _importuj_odt(self, sciezka, _=None):
        try:
            dokument = odf_wczytaj(sciezka)
            teksty = [teletype.extractText(p) for p in dokument.getElementsByType(text.P) if teletype.extractText(p).strip()]
            return True, '\n'.join(teksty), None
        except Exception as e:
            return False, "", f"Błąd ODT: {str(e)}"

    def _importuj_rtf(self, sciezka, kodowanie):
        try:
            with open(sciezka, 'r', encoding=kodowanie) as f:
                tekst = rtf_to_text(f.read())
            return True, tekst, None
        except Exception as e:
            return False, "", f"Błąd RTF: {str(e)}"

    def _importuj_braille(self, sciezka, kodowanie):
        try:
            with open(sciezka, 'r', encoding=kodowanie, errors='ignore') as f:
                zawartosc = f.read()
            if not zawartosc.strip():
                return False, "", "Plik Braille jest pusty"
            return True, zawartosc, None
        except Exception as e:
            return False, "", f"Błąd Braille: {str(e)}"

    def eksportuj_plik(self, sciezka_docelowa: str) -> tuple[bool, str]:
        if not self.aktualny_dokument:
            return False, "Brak dokumentu do eksportu"

        rozszerzenie = Path(sciezka_docelowa).suffix.lower()
        if rozszerzenie not in self.OBSLUGIWANE_EKSPORTY:
            return False, f"Format {rozszerzenie} nieobsługiwany przy eksporcie"

        tekst = self.aktualny_dokument.pobierz_zawartosc()

        try:
            if rozszerzenie == '.pdf':
                c = canvas.Canvas(sciezka_docelowa, pagesize=A4)
                c.setFont("Helvetica", 12)
                t = c.beginText(50, 800)
                for linia in tekst.split('\n'):
                    t.textLine(linia)
                c.drawText(t)
                c.save()
            elif rozszerzenie in ['.docx', '.odt', '.rtf']:
                format_doc = rozszerzenie.replace('.', '')
                pypandoc.convert_text(tekst, format_doc, format='md', outputfile=sciezka_docelowa, extra_args=['--standalone'])
            else:
                with open(sciezka_docelowa, 'w', encoding='utf-8') as f:
                    f.write(tekst)

            return True, None
        except Exception as e:
            return False, f"Błąd zapisu: {str(e)}"

    def pobierz_zawartosc(self) -> str:
        return self.aktualny_dokument.pobierz_zawartosc() if self.aktualny_dokument else None

    def ustaw_zawartosc(self, zawartosc: str):
        if self.aktualny_dokument:
            self.aktualny_dokument.ustaw_zawartosc(zawartosc)

    def cofnij(self) -> bool:
        return self.aktualny_dokument.cofnij() if self.aktualny_dokument else False

    def ponow(self) -> bool:
        return self.aktualny_dokument.ponow() if self.aktualny_dokument else False

    def pobierz_statystyki(self) -> dict:
        return self.aktualny_dokument.pobierz_statystyki() if self.aktualny_dokument else None

    def pobierz_metadane(self) -> dict:
        return self.aktualny_dokument.metadane.copy() if self.aktualny_dokument else None

    def utworz_nowy_dokument(self, zawartosc: str = ""):
        self.aktualny_dokument = Dokument(zawartosc=zawartosc)

    def czy_ma_dokument(self) -> bool:
        return self.aktualny_dokument is not None
